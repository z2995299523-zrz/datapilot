"""
DataPilot Streamlit WebUI — 数据需求分析全链路可视化

用法: streamlit run ui/app.py
"""
import sys
from pathlib import Path

# Streamlit 运行时 sys.path 只包含 ui/ 目录，需要加项目根目录才能导入 config 等模块
sys.path.insert(0, str(Path(__file__).parent.parent))

# 🔧 预加载 BGE embedding 模型（必须在 langchain_openai 之前）
#   langchain_openai 的 httpx 线程初始化会与 PyTorch CUDA 冲突，
#   导致 SentenceTransformer 在 CUDA 设备上加载时 segfault (exit 139)
from embedding import get_embedding_model
get_embedding_model()

import streamlit as st
import os

st.set_page_config(
    page_title="DataPilot — 需求分析引擎",
    page_icon="📊",
    layout="wide",
)


def _parse_requirement_file(uploaded_file) -> str:
    """从上传的需求文档中提取文本。

    支持格式:
        .txt  — 纯文本
        .md   — Markdown（按纯文本读取）
        .docx — Word 文档（提取段落文本）
    """
    import io
    name = uploaded_file.name.lower()

    if name.endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    else:
        # .txt / .md
        return uploaded_file.getvalue().decode("utf-8")

# --- Sidebar ---
st.sidebar.title("📊 DataPilot")
st.sidebar.caption("需求→SQL→测试→修复 全链路引擎")

page = st.sidebar.radio(
    "导航",
    ["📚 数据字典管理", "🔍 需求分析", "🔧 修复闭环"],
)

# --- LangSmith status indicator ---
langsmith_configured = bool(
    os.getenv("LANGCHAIN_API_KEY")
    and os.getenv("LANGCHAIN_API_KEY") != "your-langsmith-api-key"
)
if langsmith_configured:
    st.sidebar.success("🟢 LangSmith 已连接")
else:
    st.sidebar.warning("🔴 LangSmith 未配置")

st.sidebar.divider()
st.sidebar.caption("Phase 4 | 210 tests")

# --- Page routing ---
if page == "📚 数据字典管理":
    st.title("📚 数据字典管理")
    st.markdown("上传数据字典文件（CSV 或 Excel），构建 ChromaDB 向量索引。")

    # File upload
    uploaded_file = st.file_uploader(
        "选择数据字典文件",
        type=["csv", "xlsx"],
        help="支持 .csv 和 .xlsx 格式，需包含列: layer, table_name, column_name, column_type, column_comment",
    )

    if uploaded_file:
        # Save to temp file
        import tempfile
        suffix = ".xlsx" if uploaded_file.name.endswith(".xlsx") else ".csv"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        # Preview table
        st.subheader("📋 文件预览")
        import pandas as pd
        if tmp_path.endswith(".xlsx"):
            df = pd.read_excel(tmp_path)
        else:
            df = pd.read_csv(tmp_path)
        st.dataframe(df.head(20), use_container_width=True)
        st.caption(f"共 {len(df)} 行")

        # Check required columns
        required = ["layer", "table_name", "column_name"]
        missing = [c for c in required if c not in df.columns]
        if missing:
            st.error(f"❌ 缺少必要列: {', '.join(missing)}。请检查文件格式。")
        else:
            detected_layers = df['layer'].dropna().unique()
            st.success(f"✅ 列结构正确。检测到 {len(detected_layers)} 个数据层: {', '.join(detected_layers)}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("🔨 构建索引", type="primary", use_container_width=True):
                    with st.spinner("正在构建 ChromaDB 索引..."):
                        try:
                            from dictionary.loader import load_dictionary
                            from dictionary.indexer import build_index
                            data_dict = load_dictionary(tmp_path)
                            collection = build_index(data_dict, reset=True)
                            st.session_state["index_ready"] = True
                            st.session_state["index_count"] = collection.count()
                            st.success(f"✅ 索引构建完成！共 {collection.count()} 条记录")
                        except Exception as e:
                            st.error(f"❌ 索引构建失败: {e}")
            with col2:
                if st.button("🗑 重建索引", use_container_width=True):
                    with st.spinner("正在重建索引..."):
                        try:
                            from dictionary.loader import load_dictionary
                            from dictionary.indexer import build_index
                            data_dict = load_dictionary(tmp_path)
                            collection = build_index(data_dict, reset=True)
                            st.session_state["index_ready"] = True
                            st.session_state["index_count"] = collection.count()
                            st.success(f"✅ 索引重建完成！共 {collection.count()} 条记录")
                        except Exception as e:
                            st.error(f"❌ 重建失败: {e}")

        # Cleanup temp file
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    # Index status section
    st.divider()
    st.subheader("📊 索引状态")
    from config import CHROMA_DIR, CHROMA_COLLECTION
    try:
        from chromadb import PersistentClient
        from chromadb.config import Settings as ChromaSettings
        client = PersistentClient(
            path=str(CHROMA_DIR),
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        collections = client.list_collections()
        existing = [c for c in collections if c.name == CHROMA_COLLECTION]
        if existing:
            count = client.get_collection(CHROMA_COLLECTION).count()
            st.success(f"✅ 索引已就绪 — 集合 `{CHROMA_COLLECTION}` 包含 {count} 条记录")
            st.session_state["index_ready"] = True
        else:
            st.warning("⚠ 尚未构建索引。请上传数据字典文件并点击「构建索引」。")
            st.session_state["index_ready"] = False
    except Exception as e:
        st.error(f"❌ 无法连接 ChromaDB: {e}")
        st.session_state["index_ready"] = False
elif page == "🔍 需求分析":
    st.title("🔍 需求分析")
    st.markdown("输入业务需求文档，自动完成概念提取 → 分层检索 → 伪代码生成 → SQL 生成。")

    # Prerequisite check
    index_ready = st.session_state.get("index_ready", False)
    if not index_ready:
        st.warning("⚠ 请先在「数据字典管理」页面构建索引。")
        st.stop()

    # Input
    col1, col2 = st.columns([3, 1])
    with col1:
        requirement_text = st.text_area(
            "业务需求文档",
            height=200,
            placeholder="例如：统计近6个月各渠道的活跃客户数及交易金额，按渠道类型分组展示...",
            key="req_text",
        )
    with col2:
        uploaded_req = st.file_uploader(
            "或上传需求文档", type=["txt", "md", "docx"], key="req_upload",
            help="支持 .txt / .md / .docx 格式",
        )
        if uploaded_req:
            requirement_text = _parse_requirement_file(uploaded_req)
            st.text_area("已加载文件", requirement_text, height=200, disabled=True)

    generate_sql = st.checkbox("生成 SQL 脚本", value=True)

    if st.button("🚀 开始分析", type="primary", disabled=not requirement_text.strip()):
        from config import CHROMA_COLLECTION, CHROMA_DIR
        from chromadb import PersistentClient
        from chromadb.config import Settings as ChromaSettings

        client = PersistentClient(
            path=str(CHROMA_DIR),
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        collection = client.get_collection(CHROMA_COLLECTION)

        # Step 1: Concept extraction
        with st.spinner("步骤 1/3: 提取业务概念..."):
            from extractor.concept import extract_concepts
            extraction = extract_concepts(requirement_text)

        with st.expander(f"📊 步骤 1/3: 提取到 {len(extraction.concepts)} 个业务概念", expanded=True):
            for c in extraction.concepts:
                candidates = f" ({', '.join(c.candidates)})" if c.candidates else ""
                st.markdown(f"- **[{c.type.value}] {c.concept}**{candidates}")
                if c.qualifier:
                    st.caption(f"  限定: `{c.qualifier}`")

        # Step 2: Retrieval
        with st.spinner("步骤 2/3: 分层检索..."):
            from retrieval.engine import search
            result = search(extraction.concepts, collection)

        matched_count = sum(1 for m in result.matches if m.matched)
        with st.expander(f"🔍 步骤 2/3: 匹配 {matched_count}/{len(result.matches)}，未匹配 {len(result.unmatched_concepts)}", expanded=True):
            for m in result.matches:
                if m.matched:
                    codes_info = ""
                    for col in m.columns:
                        if col.code_values:
                            codes = ", ".join(f"{cv.value}={cv.meaning}" for cv in col.code_values[:6])
                            codes_info += f"  - {col.name}: {codes}\n"
                    st.success(f"**[{m.layer.value}层] {m.table_name}** (score={m.score:.2f})")
                    st.caption(m.table_comment)
                    if codes_info:
                        st.code(codes_info.strip(), language=None)
                else:
                    st.warning(f"**[未匹配] {m.concept}** — {m.message}")

            if result.unmatched_concepts:
                st.error(f"待确认: {', '.join(result.unmatched_concepts)}")

        # Step 3: Pseudocode generation
        with st.spinner("步骤 3/3: 生成分析伪代码..."):
            from generator.pseudocode import generate
            pseudocode = generate(requirement_text, result, extraction.concepts)

        with st.expander(f"📝 步骤 3/3: {pseudocode.title}", expanded=True):
            for step in pseudocode.steps:
                st.markdown(f"**步骤 {step.step_number}: {step.description}**")
                details = []
                if step.source_table:
                    details.append(f"- 源表: `{step.source_table}`")
                for cond in step.conditions:
                    details.append(f"- 条件: `{cond}`")
                for join in step.joins:
                    details.append(f"- 关联: `{join}`")
                for agg in step.aggregations:
                    details.append(f"- 聚合: `{agg}`")
                if step.output:
                    details.append(f"- 输出: {step.output}")
                st.markdown("\n".join(details))

            if pseudocode.todo_items:
                st.warning("⚠ 待确认:\n" + "\n".join(f"- {t}" for t in pseudocode.todo_items))
            if pseudocode.notes:
                st.info("📝 备注:\n" + "\n".join(f"- {n}" for n in pseudocode.notes))

        # Step 4: SQL generation (optional)
        if generate_sql:
            with st.spinner("生成 SQL 脚本..."):
                from dictionary.loader import load_dictionary
                from generator.script import generate_sql as gen_sql
                from pathlib import Path
                demo_path = Path(__file__).parent.parent / "demo" / "data_dict.csv"
                data_dict = load_dictionary(str(demo_path))
                tables = {t.table_name: t for t in data_dict.tables}
                sql = gen_sql(pseudocode, tables)

            with st.expander("💾 生成 SQL", expanded=True):
                st.code(sql, language="sql")
                st.download_button(
                    "📥 下载 SQL",
                    sql,
                    file_name="analysis.sql",
                    mime="text/plain",
                )

        # Step 5: Expected comparison (optional)
        st.divider()
        with st.expander("📊 预期结果比对（可选）", expanded=False):
            st.markdown("""
            上传预期数据 CSV，与生成的 SQL 在数据库中执行的结果逐行逐列比对。
            适用于：旧系统迁移验证、口径一致性核对、回归测试。
            """)

            expected_file = st.file_uploader(
                "上传预期数据 CSV",
                type=["csv"],
                key="expected_upload_analyze",
                help="CSV 需包含与 SQL 输出列对应的数据，用于比对验证",
            )

            if expected_file is not None:
                import pandas as pd
                try:
                    expected_df = pd.read_csv(expected_file)
                except Exception as e:
                    st.error(f"❌ CSV 解析失败: {e}")
                    st.stop()

                st.caption(f"📋 预期数据：{len(expected_df)} 行 × {len(expected_df.columns)} 列")
                st.dataframe(expected_df.head(10), use_container_width=True)

                # Need DB connection to execute SQL
                db_conn_str = st.text_input(
                    "数据库连接字符串（用于执行生成的 SQL）",
                    placeholder="sqlite:///test.db",
                    key="db_conn_analyze",
                )

                # Need generated SQL (only available if sql generation was enabled)
                if not generate_sql:
                    st.warning("⚠ 请先勾选「生成 SQL 脚本」并重新分析。")
                elif db_conn_str:
                    if st.button("🔬 执行 SQL 并比对", type="primary", key="run_compare_analyze"):
                        import tempfile
                        from sqlalchemy import create_engine, text

                        try:
                            # Save expected CSV to temp file (compare_with_expected needs path)
                            with tempfile.NamedTemporaryFile(
                                delete=False, suffix=".csv", mode="w", encoding="utf-8"
                            ) as tmp:
                                expected_df.to_csv(tmp.name, index=False)
                                tmp_path = tmp.name

                            # Execute generated SQL
                            with st.spinner("🔄 执行 SQL..."):
                                engine = create_engine(db_conn_str)
                                with engine.connect() as conn:
                                    actual_df = pd.read_sql_query(text(sql), conn)

                            st.caption(f"📊 执行结果：{len(actual_df)} 行 × {len(actual_df.columns)} 列")
                            st.dataframe(actual_df.head(10), use_container_width=True)

                            # Compare
                            with st.spinner("🔬 逐行逐列比对..."):
                                from testing.expected_compare import compare_with_expected
                                report = compare_with_expected(actual_df, tmp_path)

                            # Display report
                            st.divider()
                            if report.overall_passed:
                                st.success(f"✅ 完全匹配！预期 {report.total_expected} 行 = 实际 {report.total_actual} 行")
                            else:
                                st.error(f"❌ 发现差异：{report.summary}")

                            # Detail breakdown
                            col_a, col_b, col_c = st.columns(3)
                            with col_a:
                                st.metric("匹配行", report.match_count)
                            with col_b:
                                st.metric("差异数", report.mismatch_count)
                            with col_c:
                                status = "✅ 通过" if report.overall_passed else "❌ 失败"
                                st.metric("结果", status)

                            if report.missing_in_actual:
                                with st.expander(f"🔴 缺失行（{len(report.missing_in_actual)}）"):
                                    st.write(report.missing_in_actual[:20])
                            if report.extra_in_actual:
                                with st.expander(f"🟡 多余行（{len(report.extra_in_actual)}）"):
                                    st.write(report.extra_in_actual[:20])
                            if report.value_diffs:
                                with st.expander(f"🔵 数值偏差（{len(report.value_diffs)}）"):
                                    diffs_data = [
                                        {"键": d.key_values, "列": d.column,
                                         "预期值": d.expected_value, "实际值": d.actual_value,
                                         "偏差%": f"{d.diff_percent*100:.2f}%"}
                                        for d in report.value_diffs[:20]
                                    ]
                                    st.dataframe(pd.DataFrame(diffs_data), use_container_width=True)

                            # Cleanup temp file
                            try:
                                os.unlink(tmp_path)
                            except Exception:
                                pass

                        except Exception as e:
                            st.error(f"❌ 执行失败: {e}")

elif page == "🔧 修复闭环":
    st.title("🔧 修复闭环")
    st.markdown("输入 SQL 查询，运行 L1 数据质量 + L2 逻辑比对 + L3 诊断，自动修复并重测。")

    # Input
    original_sql = st.text_area(
        "输入需要测试的 SQL",
        height=150,
        placeholder="SELECT cust_id, channel_type, COUNT(*) as cnt\nFROM dm_customer_active\nGROUP BY cust_id, channel_type",
    )
    col_req1, col_req2 = st.columns([3, 1])
    with col_req1:
        requirement_text = st.text_area(
            "原始需求文档（可选，用于提供业务上下文）",
            height=100,
            placeholder="在此粘贴需求文档...",
            key="req_text_reconcile",
        )
    with col_req2:
        uploaded_req_reconcile = st.file_uploader(
            "或上传文档", type=["txt", "md", "docx"], key="req_upload_reconcile",
            help="支持 .txt / .md / .docx",
        )
        if uploaded_req_reconcile:
            requirement_text = _parse_requirement_file(uploaded_req_reconcile)
            st.caption(f"✅ 已加载 ({uploaded_req_reconcile.name})")

    col1, col2 = st.columns(2)
    with col1:
        max_loops = st.number_input("最大重试次数", min_value=1, max_value=10, value=3)
    with col2:
        db_conn_str = st.text_input(
            "数据库连接字符串（可选）",
            placeholder="sqlite:///test.db 或 留空跳过实际执行",
        )

    expected_file = st.file_uploader(
        "📊 预期结果 CSV（可选，用于比对验证）",
        type=["csv"],
        key="expected_upload_reconcile",
        help="上传预期数据 CSV，与 SQL 执行结果逐行逐列比对。旧系统迁移验证必备。",
    )

    if st.button("🔍 运行测试", type="primary", disabled=not original_sql.strip()):
        import json
        from models import ColumnInfo

        # Build column info from demo dict for demo purposes
        from pathlib import Path
        from dictionary.loader import load_dictionary
        demo_path = Path(__file__).parent.parent / "demo" / "data_dict.csv"
        data_dict = load_dictionary(str(demo_path))

        # Extract columns from the first DM table as demo columns
        dm_tables = [t for t in data_dict.tables if t.layer.value == "DM"]
        if dm_tables:
            column_infos = dm_tables[0].columns
        else:
            st.error("未找到 DM 层数据，请先构建索引。")
            st.stop()

        pk_columns = [c.name for c in column_infos if c.is_primary_key]

        st.divider()
        st.subheader("📊 测试执行")

        # Database connection
        conn = None
        if db_conn_str:
            try:
                from sqlalchemy import create_engine, text
                engine = create_engine(db_conn_str)
                conn = engine.connect()
                st.success("✅ 已连接数据库")
            except Exception as e:
                st.error(f"❌ 数据库连接失败: {e}。将跳过实际执行，仅展示流程。")

        # Run reconciliation
        if conn is not None:
            # Expected comparison (L2.5) — run before reconciliation
            expected_report = None
            if expected_file is not None:
                st.divider()
                with st.expander("📊 预期结果比对", expanded=True):
                    import pandas as pd
                    import tempfile

                    try:
                        # Load expected CSV
                        expected_df = pd.read_csv(expected_file)
                        st.caption(f"📋 预期数据：{len(expected_df)} 行 × {len(expected_df.columns)} 列")

                        # Execute SQL to get actual results
                        with st.spinner("🔄 执行 SQL 获取实际结果..."):
                            actual_df = pd.read_sql_query(text(original_sql), conn)

                        st.caption(f"📊 实际结果：{len(actual_df)} 行 × {len(actual_df.columns)} 列")
                        st.dataframe(actual_df.head(10), use_container_width=True)

                        # Save expected to temp file for compare_with_expected
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=".csv", mode="w", encoding="utf-8"
                        ) as tmp:
                            expected_df.to_csv(tmp.name, index=False)
                            tmp_path = tmp.name

                        with st.spinner("🔬 逐行逐列比对..."):
                            from testing.expected_compare import compare_with_expected
                            expected_report = compare_with_expected(actual_df, tmp_path)

                        # Display comparison report
                        if expected_report.overall_passed:
                            st.success(
                                f"✅ 完全匹配！预期 {expected_report.total_expected} 行 = "
                                f"实际 {expected_report.total_actual} 行"
                            )
                        else:
                            st.error(f"❌ 发现差异：{expected_report.summary}")

                        col_a, col_b, col_c = st.columns(3)
                        with col_a:
                            st.metric("匹配行", expected_report.match_count)
                        with col_b:
                            st.metric("差异数", expected_report.mismatch_count)
                        with col_c:
                            status = "✅ 通过" if expected_report.overall_passed else "❌ 失败"
                            st.metric("结果", status)

                        if expected_report.missing_in_actual:
                            with st.expander(f"🔴 缺失行（{len(expected_report.missing_in_actual)}）"):
                                st.write(expected_report.missing_in_actual[:20])
                        if expected_report.extra_in_actual:
                            with st.expander(f"🟡 多余行（{len(expected_report.extra_in_actual)}）"):
                                st.write(expected_report.extra_in_actual[:20])
                        if expected_report.value_diffs:
                            with st.expander(f"🔵 数值偏差（{len(expected_report.value_diffs)}）"):
                                diffs_data = [
                                    {
                                        "键": d.key_values, "列": d.column,
                                        "预期值": d.expected_value, "实际值": d.actual_value,
                                        "偏差%": f"{d.diff_percent*100:.2f}%",
                                    }
                                    for d in expected_report.value_diffs[:20]
                                ]
                                st.dataframe(pd.DataFrame(diffs_data), use_container_width=True)

                        # Cleanup
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass

                    except Exception as e:
                        st.error(f"❌ 预期比对失败: {e}")
                        expected_report = None

                st.divider()

            # Real execution
            try:
                from reconciliation.graph import run_reconciliation
                with st.spinner("正在执行测试与修复闭环..."):
                    final_state = run_reconciliation(
                        conn=conn,
                        original_sql=original_sql,
                        column_infos=column_infos,
                        requirement_text=requirement_text,
                        pk_columns=pk_columns,
                        expected_report_json=(
                            expected_report.model_dump_json()
                            if expected_report else ""
                        ),
                        max_loops=max_loops,
                    )

                # Display results by round
                fix_history = json.loads(final_state.get("fix_history_json", "[]"))
                for i, fix_entry in enumerate(fix_history):
                    round_num = i + 1
                    with st.expander(f"🔄 第 {round_num} 轮", expanded=(i == len(fix_history) - 1)):
                        st.json(fix_entry)

                if final_state.get("status") == "passed":
                    st.success(f"🎉 全部通过！共执行 {len(fix_history)} 轮修复")
                else:
                    st.warning(f"⚠ 测试未完全通过。状态: {final_state.get('status')}")
                    if final_state.get("error_message"):
                        st.error(final_state["error_message"])

            except Exception as e:
                st.error(f"❌ 执行失败: {e}")
            finally:
                conn.close()
        else:
            # Dry-run: show what would happen
            st.info("ℹ 未连接数据库，展示诊断流程。请连接数据库以执行实际测试。")

            with st.expander("🔍 模拟运行 — L1 基础质量检查", expanded=True):
                st.markdown("""
                将执行的检查项（基于数据字典元数据）:
                - ✅ 主键唯一性: `GROUP BY {pk} HAVING COUNT(*) > 1`
                - ✅ 空值率: 对每个字段检查 NULL 比例
                - ✅ 字段超长: 对比 varchar(N) 与 MAX(LENGTH(col))
                - ✅ 码值合法性: 对比列值与数据字典中定义的合法码值
                """)
                st.caption("实际执行时需要数据库连接。当前为流程预览。")

            with st.expander("🔧 模拟运行 — L3 诊断与修复", expanded=False):
                st.markdown("""
                如检测到失败，将触发诊断引擎:
                1. L1: 数据源检查
                2. L2: 码值映射检查
                3. L3: JOIN 逻辑检查
                4. L4: 业务口径检查
                5. L5: 概念遗漏检查

                可自动修复项: 码值替换、CAST 转换、补充 WHERE 条件
                不可自动修复: 需求口径不匹配 → 生成人工确认报告
                """)

            with st.expander("📝 结果预览", expanded=False):
                st.json({
                    "status": "dry_run",
                    "message": "此模式仅展示流程。连接数据库后可执行实际测试。",
                    "loop_count": 0,
                    "max_loops": max_loops,
                    "column_count": len(column_infos),
                    "pk_columns": pk_columns,
                })

    # Usage guide
    st.divider()
    with st.expander("💡 使用说明"):
        st.markdown("""
        **连接数据库后可执行完整的测试→诊断→修复闭环:**
        1. 输入需要测试的 SQL 查询
        2. 填写数据库连接字符串（支持 SQLite/MySQL/PostgreSQL）
        3. 点击运行测试
        4. 系统自动执行 L1/L2/L3 三层测试
        5. 失败项自动诊断并尝试修复
        6. 修复后重测，最多重试指定次数

        **不连接数据库时**展示诊断流程和可检测的问题类型。
        """)
